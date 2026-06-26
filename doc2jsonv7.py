import os                                                               # 운영체제 제어 모듈 (폴더 생성, 경로 합치기용)
import docx                                                             # MS Word(.docx) 문서를 읽기 위한 핵심 라이브러리
from docx.oxml.text.paragraph import CT_P                               # Word XML 구조 중 문단(Paragraph) 태그 타입
from docx.oxml.table import CT_Tbl                                     # Word XML 구조 중 표(Table) 태그 타입
from docx.text.paragraph import Paragraph                               # python-docx의 문단 객체 클래스
from docx.table import Table                                           # python-docx의 표 객체 클래스
import json                                                             # 파싱된 데이터를 JSON 파일로 직렬화하기 위한 모듈
import re                                                               # 문제 번호 및 보기 분리용 정규표현식 엔진

# 전역 이미지 카운터 (파일명을 고유하게 지정)
image_counter = 1                                                       # 추출되는 이미지마다 고유 번호(0001, 0002...) 부여

def iter_block_items(parent):
    """
    [함수 개요]: Word 본문/셀 내부의 요소들을 상하 배치 순서 그대로 순회 추출합니다.
    """
    for child in parent._element:                                       # 부모 요소의 XML 자식 노드들을 차례대로 탐색
        if isinstance(child, CT_P):                                     # 자식 노드가 문단(Paragraph) 구조일 경우
            yield Paragraph(child, parent)                              # 문단 객체로 감싸서 제너레이터 반환
        elif isinstance(child, CT_Tbl):                                 # 자식 노드가 표(Table) 구조일 경우
            yield Table(child, parent)                                  # 표 객체로 감싸서 제너레이터 반환

def get_paragraph_html_with_images(paragraph, doc_part, img_dir="images"):
    """
    [함수 개요]: 문단 내부를 돌며 텍스트를 수집하고, 숨겨진 이미지(Drawing/VML)를 파일로 저장 후 HTML 태그를 삽입합니다.
    """
    global image_counter                                                # 전역 이미지 카운터 변수 가져오기
    html_parts = []                                                     # 결합할 텍스트 및 HTML 태그 조각들을 담을 리스트
    
    for run in paragraph.runs:                                          # 문단 내 서식이 동일한 텍스트 조각(Run) 단위 순회
        run_text = run.text.replace('*', '')                            # 교정 기호인 별표(*) 제거 및 텍스트 가져오기
        if run_text:                                                    # 텍스트가 비어있지 않다면
            html_parts.append(run_text)                                 # 결과 리스트에 텍스트 추가
            
        drawings = run._element.xpath('.//*[local-name()="blip"]')      # 표준 DrawingML 포맷의 이미지 노드 검색
        vml_images = run._element.xpath('.//*[local-name()="imagedata"]')# 구형 VML 포맷의 이미지 노드 검색
        
        embed_ids = []                                                  # 이미지와 매핑된 고유 관계 ID(rId)를 담을 리스트
        for blip in drawings:                                           # 신형 이미지들을 돌면서
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId: embed_ids.append(rId)                               # 관계 ID가 존재하면 리스트에 추가
            
        for vml in vml_images:                                          # 구형 이미지들을 돌면서
            rId = vml.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rId: embed_ids.append(rId)                               # 관계 ID가 존재하면 리스트에 추가
            
        for embed_id in embed_ids:                                      # 수집된 모든 rId를 기반으로 이미지 추출 진행
            if embed_id in doc_part.related_parts:                      # 문서 내에 실제 미디어 바이너리가 존재하는지 확인
                img_part = doc_part.related_parts[embed_id]             # rId와 연결된 이미지 파트 객체 가져오기
                ext = img_part.content_type.split('/')[-1]              # content-type에서 확장자(jpeg, png 등) 추출
                if ext == 'jpeg': ext = 'jpg'                           # jpeg 확장자는 표준 jpg로 통일
                if ext == 'x-wmf': ext = 'png'                          # wmf 웹 비호환 확장자는 png로 변환 대피
                
                os.makedirs(img_dir, exist_ok=True)                     # 이미지를 저장할 폴더가 없으면 자동 생성
                img_filename = f"img_{image_counter:04d}.{ext}"         # 패딩을 채운 고유 파일명 생성 (예: img_0001.jpg)
                img_path = os.path.join(img_dir, img_filename)          # 디렉토리와 파일명을 결합한 전체 물리 경로 설정
                
                with open(img_path, 'wb') as f:                         # 바이너리 쓰기 모드로 이미지 파일 오픈
                    f.write(img_part.blob)                              # 이미지 데이터 원본 바이너리를 파일로 기록
                    
                # 앱 UI 테마에 맞춤화된 둥근 모서리 및 그림자 속성을 가진 웹용 반응형 <img> 태그 생성
                img_tag = f'\n<img src="{img_dir}/{img_filename}" alt="기출문제 첨부 이미지" style="max-width:100%; height:auto; border-radius:var(--border-radius-sm); margin:12px 0; box-shadow:var(--shadow-sm);">\n'
                html_parts.append(img_tag)                              # 완성된 이미지 태그를 문단 결과물에 삽입
                image_counter += 1                                      # 다음 이미지를 위해 카운터 1 증가
                
    return "".join(html_parts).strip()                                  # 문단 내 모든 텍스트와 이미지 태그를 합쳐 문자열로 반환

def get_html_from_table(table, doc_part, img_dir="images"):
    """
    [함수 개요]: 중첩된 표(Table) 객체를 파싱하여 UI 레이아웃용 HTML 테이블 또는 박스 컨테이너로 변환합니다.
    """
    if len(table.rows) == 1:                                            # 표의 행 개수가 1개일 때
        first_row = next(iter(table.rows))                              # 첫 번째 행 객체 안전하게 획득
        if len(first_row.cells) == 1:                                   # 열 개수도 1개라면 (즉, 1x1 박스 지문)
            first_cell = next(iter(first_row.cells))                    # 지문 전용 셀 객체 선택
            
            cell_html_parts = []                                        # 셀 내부 요소들의 HTML 조각 리스트
            for p in first_cell.paragraphs:                             # 셀 안의 모든 문단을 하나씩 순회
                p_html = get_paragraph_html_with_images(p, doc_part, img_dir)
                if p_html: cell_html_parts.append(p_html)               # 문단 내 텍스트/이미지를 추출하여 리스트에 축적
                
            cell_text = '<br>'.join(cell_html_parts).strip().replace('\n', '<br>') # 줄바꿈 문자를 HTML <br> 태그로 교체
            # 글라스모피즘 스타일이 적용된 강조용 지문 박스 HTML 반환
            return f'\n<div style="margin-top:16px; padding:16px; border:1px solid var(--glass-border); border-radius:var(--border-radius-sm); background:rgba(255,255,255,0.03); line-height:1.6;">{cell_text}</div>\n'
    
    html = '\n<table class="nested-table">'                             # 1x1이 아닌 일반 데이터 표는 표준 테이블 태그 시작
    for i, row in enumerate(table.rows):                                # 표의 행들을 인덱스와 함께 순회
        html += '<tr>'                                                  # 행 시작 태그 추가
        added_cells = set()                                             # 가로 병합 셀 중복 읽기 방지용 세트
        for cell in row.cells:                                          # 행 내부의 셀들을 순회
            if cell in added_cells:                                     # 이미 읽은 병합된 셀이라면
                continue                                                # 무시하고 패스
            added_cells.add(cell)                                       # 처음 보는 셀은 세트에 등록하여 중복 차단
            
            cell_html_parts = []                                        # 셀 내용물 축적용 리스트
            for p in cell.paragraphs:                                   # 셀 내부 문단 순회
                p_html = get_paragraph_html_with_images(p, doc_part, img_dir)
                if p_html: cell_html_parts.append(p_html)               # 이미지 포함 텍스트 추출 후 저장
                
            cell_text = '<br>'.join(cell_html_parts).strip().replace('\n', '<br>') # 셀 내부 개행 처리
            tag = 'th' if i == 0 else 'td'                              # 첫 번째 행은 제목 행(th), 나머지는 데이터 행(td)
            html += f'<{tag}>{cell_text}</{tag}>'                       # 셀 태그 조립
        html += '</tr>'                                                 # 행 마감
    html += '</table>\n'                                                # 테이블 마감
    return html                                                         # 완성된 테이블 HTML 반환

def parse_question_block(full_text):
    """
    [함수 개요]: 결합된 하나의 문제 블록 텍스트를 분석하여 번호, 지문, 보기, 해설, 정답을 딕셔너리로 추출합니다.
    """
    # 정규식을 이용해 시작 부분의 '숫자 + 지문' 영역 스캔 (① 기호나 '정답' 글자가 나오기 전까지가 전부 지문)
    q_match = re.search(r'^(\d+)\s+(.*?)(?=\n①|\n정답|$)', full_text, re.DOTALL)
    if not q_match:                                                     # 번호 패턴을 찾을 수 없다면 (완전한 수식이나 표 해설 덩어리)
        preview = full_text.replace('\n', ' ')[:40]                      # 가독성을 위해 개행을 공백으로 바꾸고 앞 40자만 컷
        print(f"  ⚠️ [파싱 실패] 문제 형식 불일치 (사유: '번호+지문' 패턴 없음) -> \"{preview}...\"")
        return None                                                     # 파싱 실패 로그 출력 후 빈 값 반환
    
    q_num = int(q_match.group(1))                                       # 추출된 1번 그룹(문자열 숫자)을 정수형 번호로 변환
    q_text = q_match.group(2).strip()                                   # 추출된 2번 그룹(지문 텍스트)의 앞뒤 공백 제거
    
    options = []                                                        # ①~④ 번 보기를 순서대로 담을 리스트
    for marker in ['①', '②', '③', '④']:                                # 네 개의 객관식 기호를 순차적으로 검색
        # 기호부터 다음 보기 기호나 '정답' 키워드가 나오기 전까지의 텍스트를 타겟팅
        opt_match = re.search(fr'{marker}\s*(.*?)(?=\n[①②③④]|\n정답|\n|$)', full_text, re.DOTALL)
        if opt_match:                                                   # 매칭 성공 시
            options.append(opt_match.group(1).replace('\n', ' ').strip())# 내부 줄바꿈을 한 칸 공백으로 정제하여 리스트에 추가
        else:                                                           # 보기가 없다면
            options.append("")                                          # 빈 문자열로 자리 채우기 (에러 방지)
            
    ans_match = re.search(r'정답\s*([①②③④])', full_text)                 # '정답 ①'과 같은 정답 표기 패턴 탐색
    answer_num = 0                                                      # 정답 번호 초기화 (0은 정답 없음/오류 의미)
    if ans_match:                                                       # 정답 기호를 찾았다면
        ans_map = {'①': 1, '②': 2, '③': 3, '④': 4}                    # 원문자 기호를 숫자로 매핑해주는 딕셔너리
        answer_num = ans_map.get(ans_match.group(1), 0)                 # 매핑된 정수형 정답 번호 저장
    else:                                                               # 정답 기호가 유실되었을 경우
        print(f"  ⚠️ [경고] {q_num}번 문항 - 정답 기호를 찾지 못했습니다. ('정답 ①' 형태 필요)")
    
    hint_text = full_text                                               # 전체 문제 블록 복사 (해설 텍스트 추출용 원본)
    if q_match: hint_text = hint_text.replace(q_match.group(0), '')     # 원본에서 지문 영역 제거
    if ans_match: hint_text = hint_text.replace(ans_match.group(0), '') # 원본에서 정답 영역 제거
    for marker in ['①', '②', '③', '④']:                                # 원본에서 보기 영역들 순차 제거
        opt_match = re.search(fr'{marker}\s*(.*?)(?=\n[①②③④]|\n정답|\n|$)', full_text, re.DOTALL)
        if opt_match:
            hint_text = hint_text.replace(opt_match.group(0), '')       # 매칭된 보기 텍스트 전부 소거
    
    hint_text = hint_text.strip()                                       # 남은 해설(Hint) 텍스트의 불필요한 공백 제거
    
    if not any(options):                                                # [가짜 문제 필터링]: 보기 ①~④가 단 하나도 검출되지 않았다면
        preview = full_text.replace('\n', ' ')[:40]                      # 미리보기용 텍스트 생성
        print(f"  ❌ [누락/폐기] {q_num}번 문항 폐기됨 (사유: 보기 ①~④ 없음, 병합 셀 해설/가짜 문제로 간주) -> \"{preview}...\"")
        return None                                                     # 병합 셀이 찢어지며 만들어낸 허상 문제이므로 저장하지 않고 폐기
        
    return {                                                            # 정상 변환된 최종 문제 객체 딕셔너리 리턴
        "num": q_num,
        "question": q_text,
        "options": options,
        "hint": hint_text,
        "answer": answer_num
    }

def format_question_ranges(nums):
    """
    [함수 개요]: 리스트 속 문제 번호들을 분석하여 '1~30, 32~60' 형태로 압축합니다. (V7.2 TypeError 수정본)
    """
    if not nums: return ""                                              # 번호가 존재하지 않는 비정상 회차면 빈 값 리턴
    nums = sorted(nums)                                                 # 무작위로 수집되었을 수 있으니 번호 오름차순 정렬
    ranges = []                                                         # 압축 결과를 저장할 리스트
    start = nums[0]                                                     # [버그수정]: 연속 수열의 시작점을 리스트 첫 값으로 지정
    end = nums[0]                                                       # 연속 수열의 끝점을 리스트 첫 값으로 지정
    
    for n in nums[1:]:                                                  # 두 번째 숫자부터 끝까지 반복 스캔
        if n == end + 1:                                                # 현재 숫자가 직전 숫자와 연속되어 이어진다면
            end = n                                                     # 연속 범위의 끝점을 현재 숫자로 갱신
        else:                                                           # 수열이 중간에 끊겼다면 (예: 30 다음 32)
            ranges.append(f"{start}~{end}" if start != end else str(start)) # 지금까지의 범위를 저장 (단일 값이면 단일 숫자만)
            start = n                                                   # 새로운 수열의 시작점 리셋
            end = n                                                     # 새로운 수열의 끝점 리셋
    ranges.append(f"{start}~{end}" if start != end else str(start))     # 루프가 끝나고 마지막 남은 연속 수열 범위 추가 저장
    return ", ".join(ranges)                                            # 리스트 원소들을 콤마와 공백으로 연결하여 문자열 반환

def parse_docx_to_json(docx_file, output_json):
    """
    [함수 개요]: 워드 문서의 최상위 표를 해석하고 문맥을 결합하여 파일화 및 리포트를 수행하는 메인 코어 엔진입니다.
    """
    doc = docx.Document(docx_file)                                      # 입력을 받은 워드 파일 객체 로드
    all_parsed_questions = []                                           # 전역에서 파싱 완료된 문제들을 모아놓는 보관소
    
    print("⏳ 문서 분석 및 이미지 추출을 시작합니다. 잠시만 기다려주세요...\n")
    print("-----------------------------------------------------------------")
    print("🛠️ [실시간 파싱 에러/누락 의심 로그]")
    
    for table in doc.tables:                                            # 문서에 존재하는 대형 표(회차별 문제 테이블) 순회
        added_outer_cells = set()                                       # 대형 표 내부의 세로 병합 중복 스캔 방지 메모리용 세트
        current_q_block = ""                                            # 문제 문맥 결합용 버퍼 문자열 초기화
        
        for row in table.rows:                                          # 표의 행 단위 순회 시작
            cells_content = []                                          # 행에 속한 실제 셀 텍스트들을 모을 리스트
            for cell in row.cells:                                      # 행 내부의 셀 순회
                if cell in added_outer_cells:                           # 이미 처리 완료된 병합 메모리에 기록된 셀이라면
                    continue                                            # 중복 텍스트 획득 방지를 위해 패스
                added_outer_cells.add(cell)                             # 새 셀은 중복 방지 세트에 고유 주소 등록
                
                cell_html_parts = []                                    # 단일 셀 안의 요소 파싱 결과를 축적할 리스트
                for block in iter_block_items(cell):                    # 실제 본문 순서대로 자식 노드 스캔
                    if isinstance(block, Paragraph):                    # 자식이 문단이면
                        txt = get_paragraph_html_with_images(block, doc.part, "images") # 이미지 처리 및 텍스트 획득
                        if txt:
                            cell_html_parts.append(txt)                 # 유효 텍스트만 리스트에 추가
                    elif isinstance(block, Table):                      # 자식이 표(중첩 표)라면
                        cell_html_parts.append(get_html_from_table(block, doc.part, "images")) # 표 디자인 유지 HTML 변환
                
                if cell_html_parts:                                     # 셀 내부 파싱 결과가 있다면
                    cells_content.append('\n'.join(cell_html_parts))    # 내용물들을 줄바꿈으로 연결하여 결합
            
            if not cells_content:                                       # 행에 수집된 텍스트가 아예 없다면 다음 행으로
                continue
                
            row_text = '\n'.join(cells_content)                         # 셀들의 최종 결과물을 하나의 텍스트 라인으로 결합
            
            if re.match(r'^\d+\s+', row_text):                          # [문맥 결합 핵심]: 현재 행이 새로운 '문제 번호'로 시작할 때
                if current_q_block:                                     # 기존 버퍼에 쌓여있던 문제 덩어리가 존재한다면
                    q_obj = parse_question_block(current_q_block)       # 정밀 파싱 엔진으로 덩어리 전송하여 디코딩
                    if q_obj:
                        all_parsed_questions.append(q_obj)              # 정상 문제 객체인 경우 전역 배열에 추가
                current_q_block = row_text                              # 버퍼를 비우고 현재 새로운 문제 행으로 교체 시작
            else:                                                       # 새로운 번호가 아니라면 (보기, 수식, 해설 행인 경우)
                if current_q_block:                                     # 버퍼가 활성화되어 있는 상태라면
                    current_q_block += '\n' + row_text                  # 이전 문제 덩어리의 꼬리에 강제로 붙여 문맥 결합
                else:
                    current_q_block = row_text                          # 버퍼 안전 장치 처리
        
        if current_q_block:                                             # 표 하나가 끝났을 때 버퍼에 남아있는 마지막 문제 처리
            q_obj = parse_question_block(current_q_block)               # 파싱 수행
            if q_obj:
                all_parsed_questions.append(q_obj)                      # 전역 배열에 최종 저장
                
    print("-----------------------------------------------------------------")
                
    all_rounds = []                                                     # JSON 규격에 맞게 회차별(Round)로 그룹화할 컨테이너
    current_round_questions = []                                        # 현재 회차에 소속된 문제들을 임시로 모으는 배열
    
    for q in all_parsed_questions:                                      # 수집된 전체 순수 문제들을 하나씩 확인
        if q["num"] == 1:                                               # 문제 번호가 '1'번을 만나는 순간 (새로운 기출회차 시작점)
            if current_round_questions:                                 # 이전에 이미 쌓여있던 모의고사 회차가 있었다면
                all_rounds.append({                                     # 이전 회차 데이터를 패키징하여 메인 저장소에 저장
                    "subject": "에너지관리기능장",
                    "year": "",
                    "round": f"실전모의 {len(all_rounds) + 1}회",
                    "questions": current_round_questions
                })
            current_round_questions = [q]                               # 임시 배열을 완전히 비우고, 현재 1번 문제를 첫 원소로 지정
        else:                                                           # 1번 이외의 연속된 번호(2~60번 등)라면
            if current_round_questions:                                 # 임시 배열이 정상 작동 중일 때만
                current_round_questions.append(q)                       # 현재 회차 소속으로 문제 차곡차곡 추가
                
    if current_round_questions:                                         # 문서 마지막 도달 후 남아있는 잔여 회차 마감 등록
        all_rounds.append({
            "subject": "에너지관리기능장",
            "year": "",
            "round": f"실전모의 {len(all_rounds) + 1}회",
            "questions": current_round_questions
        })
    
    with open(output_json, 'w', encoding='utf-8') as f:                 # 출력용 JSON 파일을 쓰기 전용 UTF-8 인코딩으로 오픈
        json.dump(all_rounds, f, ensure_ascii=False, indent=2)          # 한글 깨짐 방지 및 가독성 들여쓰기를 적용해 파일 기록
        
    global image_counter                                                # 전역 이미지 카운터 로드
    print(f"\n🎉 V7.2 스크립트 실행 완료!")
    print(f"총 {image_counter-1}개의 이미지가 'images/' 폴더에 추출 및 저장되었습니다.")
    
    print("\n📊 [각 회차별 문제 변환 상세 보고서]")
    print("-" * 65)
    for r in all_rounds:                                                # 마감 완료된 모든 회차 객체들을 순회하며 통계 산출
        nums = [q['num'] for q in r['questions']]                       # 해당 회차에 소속된 문제들의 번호만 리스트로 추출
        ranges_str = format_question_ranges(nums)                       # 연속 번호 수열 압축 함수 적용 (예: 1~60)
        missing_count = 60 - len(nums)                                  # 표준 기출 문항 수(60개) 기준 누락 개수 계산
        
        status = f"⚠️ 누락 {missing_count}문제" if missing_count > 0 else "✅ 완벽" # 완벽 변환 혹은 누락 발생 상태값 설정
        
        print(f" {r['round']:<10} | 총 {len(nums):2d}문제 | 번호: {ranges_str:<20} | {status}") # 가독성 정렬 프린트
    print("-" * 65)
    print(f"총 {len(all_rounds)}개 회차, 전체 {len(all_parsed_questions)}문제가 처리되었습니다.") # 최종 요약 보고

if __name__ == "__main__":
    input_file = "260626_energy_ginungjang_gyojeong.docx"               # 입력용 가공 대상 워드 문서 파일명 설정
    output_file = "questions.json"                                      # 출력용 웹 앱 호환 JSON 파일명 설정
    parse_docx_to_json(input_file, output_file)                         # 메인 컨버터 엔진 구동 시작