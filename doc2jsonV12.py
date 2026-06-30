"""
===============================================================================
[CBT 기출문제 변환 엔진 V12.9 - Hybrid Refactored Parser]
- [V12.9 리팩토링 핵심]: V12.7의 상태 머신(State Machine) 델리미터 주입 기법과 
  V12.8의 물리적 2단 표 분리(|||CELL_SPLIT|||) 기법을 유기적으로 통합.
- [동작 구조]:
  1. 2단 표(Wrapper Table)를 만났을 때는 물리적으로 좌우를 가르는 '|||CELL_SPLIT|||'을 삽입.
  2. 일반 텍스트/표 포맷 문항의 경우 상태 머신이 보기 4번 직후에 '[HINT_START]'를 실시간으로 주입.
  3. 수학 연산자/단위 기호 배제 규칙(EXCLUDE_LIST)을 적용하여 수식 오인식으로 인한 문제 찢어짐 원천 차단.
===============================================================================
"""

import os
import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
import json
import re
import datetime
import sys

# Windows 콘솔 환경용 UTF-8 설정
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')


image_counter = 1
log_file_path = ""

# 🔥 궁극의 정규식 - 수식 및 단위 문자 배제를 위한 부정 룩어헤드 및 제외 목록 탑재
EXCLUDE_LIST = [
    r'\[', r'=', r':', r'/', r'\+', r'-', r'\*', r'×', r'[xX]', r',', r'\.\d'
]
exclude_regex = '|'.join(EXCLUDE_LIST)
Q_PATTERN = rf'^[^\w\d\n]*(?:문\s*제\s*|Q\s*)?(0[1-9]|[1-5][0-9]|60)(?!\d)(?!\s*(?:{exclude_regex}))'

def log_msg(msg, level="INFO", console=True):
    if console:
        if level == "ERROR" or level == "WARNING":
            print(f"\033[91m{msg}\033[0m")
        elif level == "SUCCESS":
            print(f"\033[92m{msg}\033[0m")
        elif level == "FLUSH":
            print(f"\033[96m{msg}\033[0m")
        else:
            print(msg)
    
    if log_file_path:
        now_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        try:
            with open(log_file_path, "a", encoding="utf-8") as f:
                f.write(f"[{now_time}] [{level}] {msg}\n")
        except Exception:
            pass

def get_paragraph_html_with_images(paragraph, doc_part, subject_folder):
    global image_counter
    html_parts = []
    img_local_dir = os.path.join("data", subject_folder, "images")
    img_web_path = f"data/{subject_folder}/images"
    
    p_text = paragraph.text.replace('*', '')
    if p_text:
        html_parts.append(p_text)
        
    drawings = paragraph._element.xpath('.//*[local-name()="blip"]')
    vml_images = paragraph._element.xpath('.//*[local-name()="imagedata"]')
    
    embed_ids = []
    for blip in drawings:
        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if rId: embed_ids.append(rId)
    for vml in vml_images:
        rId = vml.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        if rId: embed_ids.append(rId)
        
    for embed_id in embed_ids:
        if embed_id in doc_part.related_parts:
            img_part = doc_part.related_parts[embed_id]
            ext = img_part.content_type.split('/').pop()
            if ext == 'jpeg': ext = 'jpg'
            if ext == 'x-wmf': ext = 'png'
            
            os.makedirs(img_local_dir, exist_ok=True)
            img_filename = f"img_{image_counter:04d}.{ext}"
            img_path = os.path.join(img_local_dir, img_filename)
            
            with open(img_path, 'wb') as f:
                f.write(img_part.blob)
                
            img_tag = f'\n<img src="{img_web_path}/{img_filename}" alt="기출문제 첨부 이미지" style="max-width:100%; height:auto; border-radius:var(--border-radius-sm); margin:12px 0; box-shadow:var(--shadow-sm);">'
            html_parts.append(img_tag)
            log_msg(f"    🖼️ [이미지 추출] {img_filename} 저장 완료", "DEBUG", console=False)
            image_counter += 1
            
    return "".join(html_parts).strip()

def get_html_from_table(table, doc_part, subject_folder):
    if len(table.rows) == 1:
        first_row = table.rows[0]
        if len(first_row.cells) == 1:
            first_cell = first_row.cells[0]
            cell_html_parts = []
            for p in first_cell.paragraphs:
                p_html = get_paragraph_html_with_images(p, doc_part, subject_folder)
                if p_html: cell_html_parts.append(p_html)
            cell_text = '<br>'.join(cell_html_parts).strip().replace('\n', '<br>')
            return f'\n<div style="margin-top:16px; padding:16px; border:1px solid var(--glass-border); border-radius:var(--border-radius-sm); background:rgba(255,255,255,0.03); line-height:1.6;">{cell_text}</div>\n'

    html = '\n<table class="nested-table">'
    for i, row in enumerate(table.rows):
        html += '<tr>'
        added_cells = set()
        for cell in row.cells:
            if cell._element in added_cells: continue
            added_cells.add(cell._element)
            cell_html_parts = []
            for p in cell.paragraphs:
                p_html = get_paragraph_html_with_images(p, doc_part, subject_folder)
                if p_html: cell_html_parts.append(p_html)
            cell_text = '<br>'.join(cell_html_parts).strip().replace('\n', '<br>')
            tag = 'th' if i == 0 else 'td'
            html += f'<{tag}>{cell_text}</{tag}>'
        html += '</tr>'
    html += '</table>\n'
    return html

def is_wrapper_table(table):
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if re.search(Q_PATTERN, text, re.MULTILINE) or re.search(r'^\s*[①②③④]', text, re.MULTILINE):
                return True
    return False

def flatten_document(parent_element, doc, doc_part, subject_folder, added_cells, lines):
    for child in parent_element:
        if child.tag.endswith('}p'):
            p = Paragraph(child, parent_element)
            txt = get_paragraph_html_with_images(p, doc_part, subject_folder)
            if txt: lines.append(txt)
        elif child.tag.endswith('}tbl'):
            if child in added_cells: continue
            added_cells.add(child)
            
            table = Table(child, parent_element)
            if is_wrapper_table(table):
                for row in table.rows:
                    row_cells = row.cells
                    unique_cells = []
                    for cell in row_cells:
                        if cell._element not in added_cells:
                            added_cells.add(cell._element)
                            unique_cells.append(cell)
                    
                    # 💡 [V12.8 구조 유지]: 왼쪽(문제/보기) 셀과 오른쪽(해설/정답) 셀 사이에 명시적인 분리 표시 삽입
                    if len(unique_cells) == 2:
                        left_lines = []
                        flatten_document(unique_cells[0]._element, doc, doc_part, subject_folder, added_cells, left_lines)
                        right_lines = []
                        flatten_document(unique_cells[1]._element, doc, doc_part, subject_folder, added_cells, right_lines)
                        
                        lines.extend(left_lines)
                        lines.append("|||CELL_SPLIT|||")
                        lines.extend(right_lines)
                    else:
                        for cell in unique_cells:
                            flatten_document(cell._element, doc, doc_part, subject_folder, added_cells, lines)
                            lines.append("[CELL_BREAK]")
            else:
                html = get_html_from_table(table, doc_part, subject_folder)
                if html:
                    lines.append("[TABLE_BREAK]")
                    lines.append(html)
                    lines.append("[TABLE_BREAK]")
        else:
            try:
                if len(child) > 0:
                    flatten_document(child, doc, doc_part, subject_folder, added_cells, lines)
            except Exception:
                pass

def parse_question_block(full_text):
    full_text = full_text.strip().replace('\t', ' ').replace('\xa0', ' ')
    
    q_num = None
    q_text = ""
    options = []
    hint_text = ""
    answer_num = 0
    
    # 💡 1. 2단 표(Wrapper Table) 분리 우선 적용 (V12.8 아키텍처)
    if "|||CELL_SPLIT|||" in full_text:
        parts = full_text.split("|||CELL_SPLIT|||", 1)
        left_part = parts[0].strip()
        right_part = parts[1].strip()
        
        # 왼쪽 영역: 문제 및 보기 파싱
        q_match = re.search(Q_PATTERN + r'(.*?)(?=[\n\s]*[①②③④]|[\n\s]*정답|$)', left_part, re.DOTALL)
        if not q_match:
            preview = left_part.replace('\n', ' ')[:60]
            log_msg(f"    ❌ [형식에러] 보기 패턴을 찾을 수 없음 -> \"{preview}...\"", "ERROR", console=False)
            return None
            
        q_num = int(q_match.group(1))
        q_text = q_match.group(2).replace('*', '').replace('[CELL_BREAK]', '').replace('[TABLE_BREAK]', '').strip()
        
        options = []
        for marker in ['①', '②', '③', '④']:
            opt_match = re.search(fr'{marker}[\s]*(.*?)(?=[\n\s]*[①②③④]|[\n\s]*정답|$)', left_part, re.DOTALL)
            if opt_match:
                raw_opt = opt_match.group(1).replace('[CELL_BREAK]', '').replace('[TABLE_BREAK]', '').replace('\n', ' ').strip()
                options.append(raw_opt)
            else:
                options.append("")
                
        # 오른쪽 영역: 정답 및 힌트 파싱
        ans_match = re.search(r'정답[\s]*([①②③④])', right_part)
        if ans_match:
            ans_map = {'①': 1, '②': 2, '③': 3, '④': 4}
            answer_num = ans_map.get(ans_match.group(1), 0)
            
        hint_text = right_part
        if ans_match:
            hint_text = hint_text.replace(ans_match.group(0), '')
        hint_text = hint_text.replace("|||CELL_SPLIT|||", "\n").replace('[CELL_BREAK]', '\n').replace('[TABLE_BREAK]', '\n').strip()
        hint_text = re.sub(r'\n+', '\n', hint_text).strip()
        
    # 💡 2. 일반 텍스트 문항: V12.7 [HINT_START] 상태 머신 기반 파싱 적용
    elif '[HINT_START]' in full_text:
        parts = full_text.split('[HINT_START]', 1)
        q_and_opts = parts[0].strip()
        hint_part = parts[1].strip()
        
        q_match = re.search(Q_PATTERN + r'(.*?)(?=[\n\s]*[①②③④]|$)', q_and_opts, re.DOTALL)
        if not q_match:
            preview = q_and_opts.replace('\n', ' ')[:60]
            log_msg(f"    ❌ [형식에러] 보기 패턴을 찾을 수 없음 -> \"{preview}...\"", "ERROR", console=False)
            return None
            
        q_num = int(q_match.group(1))
        q_text = q_match.group(2).replace('*', '').replace('[CELL_BREAK]', '').replace('[TABLE_BREAK]', '').strip()
        
        options = []
        for marker in ['①', '②', '③', '④']:
            opt_match = re.search(fr'{marker}[\s]*(.*?)(?=[\n\s]*[①②③④]|$)', q_and_opts, re.DOTALL)
            if opt_match:
                raw_opt = opt_match.group(1).replace('[CELL_BREAK]', '').replace('[TABLE_BREAK]', '').replace('\n', ' ').strip()
                options.append(raw_opt)
            else:
                options.append("")
                
        ans_match = re.search(r'정답[\s]*([①②③④])', full_text)
        if ans_match:
            ans_map = {'①': 1, '②': 2, '③': 3, '④': 4}
            answer_num = ans_map.get(ans_match.group(1), 0)
            
        hint_text = hint_part
        if ans_match:
            hint_text = hint_text.replace(ans_match.group(0), '')
        hint_text = hint_text.replace('[CELL_BREAK]', '\n').replace('[TABLE_BREAK]', '\n').strip()
        hint_text = re.sub(r'\n+', '\n', hint_text).strip()
        
    else:
        # 💡 3. 폴백: 구분자가 둘 다 없는 경우 (기본 방식 복원)
        q_match = re.search(Q_PATTERN + r'(.*?)(?=[\n\s]*[①②③④]|[\n\s]*정답|$)', full_text, re.DOTALL)
        if not q_match:
            preview = full_text.replace('\n', ' ')[:60]
            log_msg(f"    ❌ [형식에러] 보기 패턴을 찾을 수 없음 -> \"{preview}...\"", "ERROR", console=False)
            return None
            
        q_num = int(q_match.group(1))
        q_text = q_match.group(2).replace('*', '').replace('[CELL_BREAK]', '').replace('[TABLE_BREAK]', '').strip()
        
        options = []
        for marker in ['①', '②', '③']:
            opt_match = re.search(fr'{marker}[\s]*(.*?)(?=[\n\s]*[①②③④]|[\n\s]*정답|$)', full_text, re.DOTALL)
            if opt_match:
                options.append(opt_match.group(1).replace('[CELL_BREAK]', '').replace('[TABLE_BREAK]', '').replace('\n', ' ').strip())
            else:
                options.append("")
                
        opt_match = re.search(fr'④[\s]*(.*?)(?=\n|정답|$)', full_text, re.DOTALL)
        if opt_match:
            options.append(opt_match.group(1).replace('[CELL_BREAK]', '').replace('[TABLE_BREAK]', '').replace('\n', ' ').strip())
        else:
            options.append("")
            
        ans_match = re.search(r'정답[\s]*([①②③④])', full_text)
        if ans_match:
            ans_map = {'①': 1, '②': 2, '③': 3, '④': 4}
            answer_num = ans_map.get(ans_match.group(1), 0)
            
        hint_text = full_text
        if q_match: hint_text = hint_text.replace(q_match.group(0), '')
        if ans_match: hint_text = hint_text.replace(ans_match.group(0), '')
        for marker in ['①', '②', '③', '④']:
            opt_m = re.search(fr'{marker}[\s]*(.*?)(?=[\n\s]*[①②③④]|[\n\s]*정답|$)', full_text, re.DOTALL)
            if opt_m:
                hint_text = hint_text.replace(opt_m.group(0), '')
        hint_text = hint_text.replace('[CELL_BREAK]', '\n').replace('[TABLE_BREAK]', '\n').strip()
        hint_text = re.sub(r'\n+', '\n', hint_text).strip()

    if not any(options):
        preview = full_text.replace('\n', ' ')[:60]
        log_msg(f"    ⚠️ [보기누락] {q_num}번 문항 폐기 -> \"{preview}...\"", "WARNING", console=False)
        return None
        
    log_msg(f"    ✅ [문항 파싱 성공] {q_num:02d}번 추출 (정답: {answer_num}) -> 지문: \"{q_text[:30]}...\"", "SUCCESS", console=True)
    
    return {
        "num": q_num,
        "question": q_text,
        "options": options,
        "hint": hint_text,
        "answer": answer_num
    }

def format_question_ranges(nums):
    if not nums: return ""
    nums_sorted = sorted(set(nums))
    ranges = []
    num_iter = iter(nums_sorted)
    start = next(num_iter)
    end = start
    
    for n in num_iter:
        if n == end + 1:
            end = n
        else:
            ranges.append(f"{start}~{end}" if start != end else str(start))
            start = n
            end = n
            
    ranges.append(f"{start}~{end}" if start != end else str(start))
    return ", ".join(ranges)

def parse_docx_to_json(docx_file, output_json, subject_folder):
    global log_file_path
    
    base_name = os.path.splitext(os.path.basename(docx_file))[0]
    now_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file_path = f"({repr(base_name)}, '.docx')_{now_str}.log.txt"
    
    with open(log_file_path, "w", encoding="utf-8") as f:
        f.write("="*70 + "\n")
        f.write(" CBT 기출문제 변환 엔진 V12.9 - Hybrid Refactored Architecture\n")
        f.write(f" [날짜/시간] : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("="*70 + "\n\n")

    doc = docx.Document(docx_file)
    
    subject_map = {
        "energy_ginungjang": "에너지관리기능장",
        "energy_sanupgisa": "에너지관리산업기사",
        "gas": "가스기능사",
        "air_conditioning": "공조기능사"
    }
    real_subject_name = subject_map.get(subject_folder, "기출문제")
    
    log_msg(f"\n⏳ [{real_subject_name}] 문서 선형화 진행 중...", "INFO", console=True)
    
    lines = []
    added_cells = set()
    flatten_document(doc.element.body, doc, doc.part, subject_folder, added_cells, lines)
    
    log_msg("🔄 [1-Pass 스마트 파싱] 문서 스캔 및 델리미터 주입 진행...", "INFO", console=True)
    
    all_rounds = []
    current_round_questions = []
    current_round_info = None
    mock_counter = 1
    current_q_block = ""
    current_state = "NONE" # 상태: NONE, QUESTION, OPT1~4, HINT
    
    for line in lines:
        clean_line = line.replace('\ufeff', '').replace('\u200b', '')
        
        # 1. 회차 헤더 감지 시 FLUSH 로직
        header_match = re.search(r'(\d{4})년\s*(\d+)회', clean_line)
        if header_match:
            # 진행 중이던 문제 블록이 있으면 먼저 닫아서 넣음
            if current_q_block:
                q_obj = parse_question_block(current_q_block)
                if q_obj: current_round_questions.append(q_obj)
                current_q_block = ""
                
            # 이전 회차 데이터가 존재하면 JSON 배열로 넘기고 메모리 리셋
            if current_round_questions:
                if not current_round_info:
                    current_round_info = {"year": "", "round": f"실전모의 {mock_counter}회"}
                    mock_counter += 1
                    
                all_rounds.append({
                    "subject": real_subject_name,
                    "year": current_round_info["year"],
                    "round": current_round_info["round"],
                    "questions": current_round_questions
                })
                log_msg(f"\n💾 [FLUSH 완료] {current_round_info['year']}년 {current_round_info['round']} - 총 {len(current_round_questions)}문항 분리 및 저장됨\n", "FLUSH", console=True)
                
            current_round_info = {
                "year": int(header_match.group(1)),
                "round": f"{header_match.group(2)}회"
            }
            current_round_questions = []
            current_state = "NONE"
            log_msg(f"📌 [신규 회차 스캔 시작] {current_round_info['year']}년 {current_round_info['round']}", "INFO", console=True)
            continue
        
        # 2. 문제 번호 감지 및 추가 (수식 및 괄호 필터링 탑재)
        q_match = re.match(Q_PATTERN, clean_line.strip())
        if q_match:
            matched_prefix = q_match.group(0)
            remaining = clean_line.strip()[len(matched_prefix):].strip()
            # 숫자 뒤에 대괄호 기호로 시작하거나(예: 10[L]), 13글자 미만의 짧은 줄은 수식으로 간주
            if not remaining.startswith('[') and len(clean_line.strip()) >= 13:
                if current_q_block:
                    q_obj = parse_question_block(current_q_block)
                    if q_obj: current_round_questions.append(q_obj)
                    
                current_q_block = line
                current_state = "QUESTION"
                q_num = q_match.group(1)
                reason = q_match.group(0).strip()
                preview = clean_line[:40].replace('\n', ' ')
                log_msg(f"  🆕 [문항 인식] 패턴: '{reason}' -> {q_num}번 문항 | 원문: \"{preview}...\"", "INFO", console=True)
                continue
        
        # 3. 문제 블록 누적 및 상태 머신에 의한 [HINT_START] 주입
        if current_q_block:
            # 힌트 상태가 아닐 때만 보기를 감지하여 [HINT_START] 삽입
            if current_state != "HINT":
                if re.match(r'^\s*①', clean_line):
                    current_state = "OPT1"
                elif re.match(r'^\s*②', clean_line):
                    current_state = "OPT2"
                elif re.match(r'^\s*③', clean_line):
                    current_state = "OPT3"
                elif re.match(r'^\s*④', clean_line):
                    current_state = "OPT4"
                elif re.match(r'^\s*정답', clean_line):
                    current_q_block += '\n[HINT_START]'
                    current_state = "HINT"
                elif current_state == "OPT4" and clean_line.strip() != "":
                    current_q_block += '\n[HINT_START]'
                    current_state = "HINT"
            
            current_q_block += '\n' + line
        else:
            if clean_line.strip() and re.search(r'^[^\w\d\n]*\d', clean_line):
                preview = clean_line[:50].replace('\n', ' ')
                log_msg(f"  🗑️ [폐기] 문제 아님(해설/수식) -> \"{preview}...\"", "WARNING", console=False)

    # 마지막 문서 끝에 남은 문제 블록 처리
    if current_q_block:
        q_obj = parse_question_block(current_q_block)
        if q_obj: current_round_questions.append(q_obj)

    # 마지막 회차 FLUSH
    if current_round_questions:
        if not current_round_info:
            current_round_info = {"year": "", "round": f"실전모의 {mock_counter}회"}
            
        all_rounds.append({
            "subject": real_subject_name,
            "year": current_round_info["year"],
            "round": current_round_info["round"],
            "questions": current_round_questions
        })
        log_msg(f"\n💾 [FLUSH 완료] {current_round_info['year']}년 {current_round_info['round']} - 총 {len(current_round_questions)}문항 분리 및 저장됨\n", "FLUSH", console=True)
    
    # JSON 파일 저장
    output_dir = os.path.dirname(output_json)
    if output_dir: os.makedirs(output_dir, exist_ok=True)

    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(all_rounds, f, ensure_ascii=False, indent=2)
        
    global image_counter
    log_msg(f"🎉 V12.9 Hybrid Refactored Architecture 실행 완료!", "SUCCESS", console=True)
    log_msg(f"상세 내역이 '{log_file_path}'에 저장되었습니다.", "INFO", console=True)
    
    log_msg("\n📊 [각 회차별 문제 변환 상세 보고서]", "INFO", console=True)
    log_msg("-" * 65, "INFO", console=True)
    
    total_parsed_questions = 0
    for r in all_rounds:
        nums = [q['num'] for q in r['questions']]
        round_question_count = len(nums)
        total_parsed_questions += round_question_count 
        ranges_str = format_question_ranges(nums)
        missing_count = 60 - round_question_count
        status = f"⚠️ 누락 {missing_count}문제" if missing_count > 0 else "✅ 완벽"
        
        round_title = f"{r['year']}년 {r['round']}" if r['year'] else r['round']
        log_msg(f" {round_title:<12} | 총 {round_question_count:2d}문제 | 번호: {ranges_str:<20} | {status}", "INFO", console=True)
        
    log_msg("-" * 65, "INFO", console=True)
    log_msg(f"🎯 [최종 결과] 총 {len(all_rounds)}개 회차, 전체 {total_parsed_questions}문항 변환 성공!", "SUCCESS", console=True)
    log_msg("-" * 65, "INFO", console=True)

if __name__ == "__main__":
    subject_name = "gas"
    input_file = "gas_CBT_2017_2025.docx"
    output_file = f"data/{subject_name}/{subject_name}_questions.json"
    
    parse_docx_to_json(input_file, output_file, subject_name)
